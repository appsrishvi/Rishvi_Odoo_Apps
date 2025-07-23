from odoo import models, fields, api
from datetime import datetime, timedelta
from collections import defaultdict
import base64
from io import BytesIO
import xlsxwriter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle, SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT


class DailyTimesheetWizard(models.TransientModel):
    _name = 'daily.timesheet.wizard'
    _description = 'Daily Timesheet Report Wizard'

    # File output
    file = fields.Binary("File")
    file_name = fields.Char("Filename")
    file_format = fields.Selection([
        ('xlsx', 'XLSX'),
        ('docx', 'DOCX'),
        ('pdf', 'PDF')
    ], string="File Format", default='xlsx')

    # Date Filters
    date_option = fields.Selection([
        ('today', 'Today'),
        ('custom', 'Custom Date Range')
    ], string="Date Filter", default='today')

    start_date = fields.Date("Start Date")
    end_date = fields.Date("End Date")

    # Project Filters
    project_id = fields.Many2one('project.project', string="Project")

    # Employee Filters
    employee_option = fields.Selection([
        ('all', 'All Employees'),
        ('custom', 'Custom Employees')
    ], string="Employee Filter", default='all')

    employee_ids = fields.Many2many('hr.employee', string="Employees")

    @api.onchange('project_id')
    def _onchange_project_id(self):
        if self.project_id:
            tasks = self.env['project.task'].search([('project_id', '=', self.project_id.id)])
            users = tasks.mapped('user_ids')
            employees = self.env['hr.employee'].search([('user_id', 'in', users.ids)])
            self.employee_ids = [(6, 0, employees.ids)]
            return {
                'domain': {
                    'employee_ids': [('id', 'in', employees.ids)]
                }
            }
        else:
            # If no project, show no employees
            self.employee_ids = [(6, 0, [])]
            return {'domain': {'employee_ids': [('id', '=', False)]}}


    @api.onchange('date_option')
    def _onchange_date_option(self):
        if self.date_option != 'custom':
            self.start_date = False
            self.end_date = False

    @api.onchange('employee_option')
    def _onchange_employee_option(self):
        if self.employee_option != 'custom':
            self.employee_ids = [(5, 0, 0)]

    def _get_grouped_timesheets(self):
        # Build domain for timesheets
        domain = []

        # Date filter
        if self.date_option == 'today':
            today = datetime.today().date()
            domain += [('date', '=', today)]
        elif self.date_option == 'custom':
            if self.start_date:
                domain.append(('date', '>=', self.start_date))
            if self.end_date:
                domain.append(('date', '<=', self.end_date))

        #Project filter
        if self.project_id:
            task_users = self.env['project.task'].search([
                ('project_id', '=', self.project_id.id)
            ]).mapped('user_ids')
            employees = task_users.mapped('employee_ids')

            if self.employee_option == 'all':
                domain.append(('employee_id', 'in', employees.ids))
            elif self.employee_option == 'custom' and self.employee_ids:
                domain.append(('employee_id', 'in', self.employee_ids.ids))
        else:
            # Employee filter
            if self.employee_option == 'custom' and self.employee_ids:
                domain.append(('employee_id', 'in', self.employee_ids.ids))

        timesheets = self.env['account.analytic.line'].search(domain)

        # Group by employee
        grouped = defaultdict(list)
        for ts in timesheets:
            key = ts.employee_id.name or 'Unknown'
            grouped[key].append(ts)

        return grouped

    def action_import(self):
        if self.file_format == 'xlsx':
            return self._export_xlsx()
        elif self.file_format == 'docx':
            return self._export_docx()
        elif self.file_format == 'pdf':
            return self._export_pdf()

    def _export_xlsx(self):
        grouped = self._get_grouped_timesheets()

        output = BytesIO()
        workbook = xlsxwriter.Workbook(output, {'in_memory': True})
        bold = workbook.add_format({'bold': True})
        header_format = workbook.add_format({'bold': True, 'bg_color': '#DCE6F1'})
        wrap_format = workbook.add_format({'text_wrap': True})
        total_format = workbook.add_format({'bold': True, 'font_color': 'blue'})

        sheet = workbook.add_worksheet('Timesheet')
        row = 0

        sheet.write(row, 0, 'Daily Timesheet Report', bold)
        row += 2

        for employee, records in grouped.items():
            sheet.write(row, 0, employee, bold)
            row += 1
            sheet.write_row(row, 0, ['Project', 'Task', 'Date', 'Hours'], header_format)
            row += 1
            total = 0
            for rec in records:
                sheet.write(row, 0, rec.project_id.name, wrap_format)
                sheet.write(row, 1, rec.name)
                sheet.write(row, 2, str(rec.date))
                sheet.write(row, 3, rec.unit_amount)
                total += rec.unit_amount
                row += 1
            sheet.write_row(row, 0, ['Total Hours', '', '', total], total_format)
            row += 2

        workbook.close()
        output.seek(0)

        self.write({
            'file': base64.b64encode(output.getvalue()),
            'file_name': 'daily_timesheet_report.xlsx',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f"/web/content/?model={self._name}&id={self.id}&field=file&filename={self.file_name}&download=true",
            'target': 'self'
        }

    def _export_docx(self):
        grouped = self._get_grouped_timesheets()

        document = Document()
        document.add_heading("Daily Timesheet Report", level=0)

        for employee, records in grouped.items():
            document.add_heading(employee, level=1)
            table = document.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Project'
            hdr_cells[1].text = 'Task'
            hdr_cells[2].text = 'Date'
            hdr_cells[3].text = 'Hours'

            total_hours = 0
            for rec in records:
                row_cells = table.add_row().cells
                row_cells[0].text = rec.project_id.name or ''
                row_cells[1].text = rec.name or ''
                row_cells[2].text = rec.date.strftime('%Y-%m-%d')
                row_cells[3].text = f'{rec.unit_amount:.2f}'
                total_hours += rec.unit_amount

                row_cells[0].width = Inches(2.5)
                row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)

            total_row = table.add_row().cells
            total_row[0].merge(total_row[1]).merge(total_row[2])
            total_row[0].text = "Total Hours"
            total_row[3].text = f'{total_hours:.2f}'
            for cell in [total_row[0], total_row[3]]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 255)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        self.write({
            'file': base64.b64encode(buffer.getvalue()),
            'file_name': 'daily_timesheet_report.docx',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f"/web/content/?model={self._name}&id={self.id}&field=file&filename={self.file_name}&download=true",
            'target': 'self'
        }

    def _export_pdf(self):
        grouped = self._get_grouped_timesheets()

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        elements.append(Paragraph("Daily Timesheet Report", styles['Title']))
        elements.append(Spacer(1, 12))

        for employee, records in grouped.items():
            elements.append(Paragraph(f"<b>Employee:</b> {employee}", styles['Heading2']))

            data = [['Project', 'Task', 'Date', 'Hours']]
            total = 0
            for rec in records:
                data.append([
                    Paragraph(rec.project_id.name or '', ParagraphStyle(name='Wrapped', wordWrap='CJK', fontSize=9)),
                    Paragraph(rec.name or '', styles['Normal']),
                    rec.date.strftime('%Y-%m-%d'),
                    f"{rec.unit_amount:.2f}"
                ])
                total += rec.unit_amount

            data.append([
                Paragraph(f'<font color="blue"><b>Total Hours</b></font>', styles['Normal']), '', '',
                Paragraph(f'<font color="blue"><b>{total:.2f}</b></font>', styles['Normal'])
            ])

            table = Table(data, colWidths=[130, 200, 100, 50])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#DCE6F1')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
                ('FONTNAME', (-1, -1), (-1, -1), 'Helvetica-Bold'),
            ]))

            elements.append(table)
            elements.append(Spacer(1, 24))

        doc.build(elements)
        buffer.seek(0)

        self.write({
            'file': base64.b64encode(buffer.getvalue()),
            'file_name': 'daily_timesheet_report.pdf',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f"/web/content/?model={self._name}&id={self.id}&field=file&filename={self.file_name}&download=true",
            'target': 'self'
        }

